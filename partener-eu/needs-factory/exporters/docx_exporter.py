from __future__ import annotations

import hashlib
import json
import re
import tempfile
import zipfile
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Mapping, Sequence, Tuple
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Mm, Pt


NEED_TAG_RE = re.compile(r"\[NEED:([^\]]+)\]")
EVIDENCE_TAG_RE = re.compile(r"\[EV:([^\]]+)\]")
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")
ZIP_EPOCH = (1980, 1, 1, 0, 0, 0)
FIXED_DOC_DATE = datetime(2000, 1, 1, 0, 0, 0)
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class ExportValidationError(ValueError):
    """Raised when a downstream export would lose canonical semantic markers."""


def _sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def sha256_file(path: Path) -> str:
    return _sha256_bytes(path.read_bytes())


def _set_document_defaults(document: Document, *, title: str) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    for style_name, size, bold in (
        ("Normal", 10.5, False),
        ("Title", 18, True),
        ("Heading 1", 15, True),
        ("Heading 2", 13, True),
        ("Heading 3", 11.5, True),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = bold

    props = document.core_properties
    props.title = title
    props.author = "Needs Factory"
    props.last_modified_by = "Needs Factory"
    props.created = FIXED_DOC_DATE
    props.modified = FIXED_DOC_DATE
    props.comments = "Compiled downstream from a validated Needs Factory narrative pack."


def _add_inline(paragraph, text: str) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            run.font.name = "Arial"
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
        else:
            run = paragraph.add_run(token)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        run.font.name = "Arial"


def _table_row(line: str) -> List[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def _is_table_separator(line: str) -> bool:
    cells = _table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def _set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.bold = bold


def markdown_to_docx(markdown: str, output_path: Path, *, metadata_title: str) -> None:
    document = Document()
    _set_document_defaults(document, title=metadata_title)

    lines = markdown.splitlines()
    index = 0
    first_h1 = True
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|") and lines[index].strip().endswith("|"):
                table_lines.append(lines[index])
                index += 1
            if len(table_lines) >= 2 and _is_table_separator(table_lines[1]):
                header = _table_row(table_lines[0])
                body = [_table_row(row) for row in table_lines[2:]]
                table = document.add_table(rows=1, cols=len(header))
                table.style = "Table Grid"
                for col, text in enumerate(header):
                    _set_cell_text(table.rows[0].cells[col], text, bold=True)
                for row_data in body:
                    row_cells = table.add_row().cells
                    for col in range(len(header)):
                        _set_cell_text(row_cells[col], row_data[col] if col < len(row_data) else "")
                continue
            for raw in table_lines:
                paragraph = document.add_paragraph()
                _add_inline(paragraph, raw)
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading_match:
            hashes, text = heading_match.groups()
            level = len(hashes)
            if level == 1 and first_h1:
                paragraph = document.add_paragraph(style="Title")
                first_h1 = False
            else:
                paragraph = document.add_heading(level=level)
            _add_inline(paragraph, text)
            index += 1
            continue

        bullet_match = re.match(r"^(\s*)-\s+(.*)$", line)
        if bullet_match:
            indent, text = bullet_match.groups()
            paragraph = document.add_paragraph(style="List Bullet")
            if len(indent) >= 2:
                paragraph.paragraph_format.left_indent = Mm(8)
            _add_inline(paragraph, text)
            index += 1
            continue

        paragraph = document.add_paragraph()
        _add_inline(paragraph, stripped)
        index += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        temp_path = Path(handle.name)
    try:
        document.save(temp_path)
        _normalize_zip(temp_path, output_path)
    finally:
        temp_path.unlink(missing_ok=True)


def _normalize_zip(source: Path, target: Path) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(source, "r") as src, zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as dst:
        for name in sorted(src.namelist()):
            data = src.read(name)
            original = src.getinfo(name)
            info = zipfile.ZipInfo(filename=name, date_time=ZIP_EPOCH)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = original.external_attr
            info.internal_attr = original.internal_attr
            info.create_system = original.create_system
            info.flag_bits = original.flag_bits & ~0x08
            dst.writestr(info, data, compress_type=zipfile.ZIP_DEFLATED, compresslevel=9)


def _docx_text(path: Path) -> str:
    with zipfile.ZipFile(path, "r") as archive:
        xml = archive.read("word/document.xml")
    root = ET.fromstring(xml)
    texts = [node.text or "" for node in root.iter(f"{{{W_NS}}}t")]
    return " ".join(texts)


def validate_docx_semantics(docx_path: Path, source_markdown: str) -> Dict[str, Any]:
    extracted = _docx_text(docx_path)
    source_need_counts = Counter(NEED_TAG_RE.findall(source_markdown))
    source_evidence_counts = Counter(EVIDENCE_TAG_RE.findall(source_markdown))
    docx_need_counts = Counter(NEED_TAG_RE.findall(extracted))
    docx_evidence_counts = Counter(EVIDENCE_TAG_RE.findall(extracted))
    failures: List[Dict[str, Any]] = []

    if source_need_counts != docx_need_counts:
        failures.append({
            "failure": "need_marker_mismatch",
            "source": dict(sorted(source_need_counts.items())),
            "docx": dict(sorted(docx_need_counts.items())),
        })
    if source_evidence_counts != docx_evidence_counts:
        failures.append({
            "failure": "evidence_marker_mismatch",
            "source": dict(sorted(source_evidence_counts.items())),
            "docx": dict(sorted(docx_evidence_counts.items())),
        })
    source_limit_count = source_markdown.count("Limită de interpretare:")
    docx_limit_count = extracted.count("Limită de interpretare:")
    if source_limit_count != docx_limit_count:
        failures.append({
            "failure": "interpretation_limit_count_mismatch",
            "source": source_limit_count,
            "docx": docx_limit_count,
        })

    return {
        "schema_version": "nf.docx_validation.v0.1",
        "valid": not failures,
        "failures": failures,
        "need_marker_count": sum(docx_need_counts.values()),
        "evidence_marker_count": sum(docx_evidence_counts.values()),
        "interpretation_limit_count": docx_limit_count,
        "document_xml_sha256": _sha256_bytes(zipfile.ZipFile(docx_path, "r").read("word/document.xml")),
    }


def _write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8", newline="\n")


def _deterministic_zip(paths: Sequence[Tuple[Path, str]], output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
        for source, arcname in sorted(paths, key=lambda item: item[1]):
            info = zipfile.ZipInfo(arcname, ZIP_EPOCH)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.create_system = 0
            archive.writestr(info, source.read_bytes(), compress_type=zipfile.ZIP_DEFLATED, compresslevel=9)


def export_final_package(
    compiled_analysis: Mapping[str, Any],
    output_dir: Path,
    *,
    basename: str = "ANALIZA_NEVOI",
) -> Dict[str, Any]:
    validation = compiled_analysis.get("validation") or {}
    if not validation.get("valid"):
        raise ExportValidationError("compiled analysis must pass narrative validation before export")

    output_dir.mkdir(parents=True, exist_ok=True)
    analysis_md = output_dir / f"{basename}.md"
    analysis_docx = output_dir / f"{basename}.docx"
    sources_md = output_dir / "SOURCE_REGISTER.md"
    sources_docx = output_dir / "SOURCE_REGISTER.docx"
    manifest_path = output_dir / "PACKAGE_MANIFEST.json"
    package_zip = output_dir / f"{basename}_PACKAGE.zip"

    markdown = str(compiled_analysis["markdown"])
    source_register = str(compiled_analysis["source_register_markdown"])
    _write_text(analysis_md, markdown)
    _write_text(sources_md, source_register)
    markdown_to_docx(markdown, analysis_docx, metadata_title=basename.replace("_", " ").title())
    markdown_to_docx(source_register, sources_docx, metadata_title="Needs Factory Source Register")

    analysis_validation = validate_docx_semantics(analysis_docx, markdown)
    source_validation = validate_docx_semantics(sources_docx, source_register)
    if not analysis_validation["valid"]:
        raise ExportValidationError(f"analysis DOCX failed semantic validation: {analysis_validation['failures']}")
    if not source_validation["valid"]:
        raise ExportValidationError(f"source register DOCX failed semantic validation: {source_validation['failures']}")

    files = {
        analysis_md.name: {"role": "compiled_analysis_markdown", "sha256": sha256_file(analysis_md)},
        analysis_docx.name: {"role": "compiled_analysis_docx", "sha256": sha256_file(analysis_docx)},
        sources_md.name: {"role": "source_register_markdown", "sha256": sha256_file(sources_md)},
        sources_docx.name: {"role": "source_register_docx", "sha256": sha256_file(sources_docx)},
    }
    manifest = {
        "schema_version": "nf.final_package.v0.1",
        "source_pack_sha256": compiled_analysis.get("source_pack_sha256"),
        "source_markdown_sha256": compiled_analysis.get("markdown_sha256"),
        "source_register_sha256": compiled_analysis.get("source_register_sha256"),
        "downstream_only": True,
        "claim_mutation_allowed": False,
        "files": files,
        "docx_validation": {
            analysis_docx.name: analysis_validation,
            sources_docx.name: source_validation,
        },
    }
    _write_text(manifest_path, json.dumps(manifest, ensure_ascii=False, indent=2, sort_keys=True) + "\n")
    manifest["files"][manifest_path.name] = {"role": "package_manifest", "sha256": sha256_file(manifest_path)}

    _deterministic_zip(
        [(analysis_md, analysis_md.name), (analysis_docx, analysis_docx.name), (sources_md, sources_md.name), (sources_docx, sources_docx.name), (manifest_path, manifest_path.name)],
        package_zip,
    )
    manifest["package_zip"] = {"path": package_zip.name, "sha256": sha256_file(package_zip)}
    return manifest
