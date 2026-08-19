from __future__ import annotations

import csv
import hashlib
import json
import tempfile
from pathlib import Path
from typing import Any, Dict, Mapping

from docx import Document
from docx.shared import Pt

from .docx_exporter import _deterministic_zip, _normalize_zip, _set_document_defaults, sha256_file


class ResearchPackError(ValueError):
    """Raised when a primary-research plan is not suitable for a collection pack."""


def _write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8", newline="\n")


def _response_instruction(response_type: str) -> str:
    if response_type in {"likert_1_5", "likert_1_5_optional"}:
        return "Bifează un singur răspuns: 1 = nivel foarte scăzut; 2 = scăzut; 3 = mediu; 4 = ridicat; 5 = foarte ridicat."
    if response_type == "yes_no":
        return "Bifează un singur răspuns: Da / Nu."
    return f"Completează răspunsul conform tipului de răspuns: {response_type}."


def render_questionnaire_markdown(plan: Mapping[str, Any]) -> str:
    questions = list(plan.get("questions") or [])
    if not questions:
        raise ResearchPackError("primary research plan has no questions")
    lines = [
        "# Chestionar pentru analiza de nevoi",
        "",
        "Acest chestionar este utilizat exclusiv pentru fundamentarea analizei de nevoi. Nu se solicită nume, CNP sau alte date de identificare directă. Fiecărui respondent i se atribuie un cod pseudonimizat.",
        "",
        "## Instrucțiuni pentru respondent",
        "",
        "Răspunde în raport cu experiența și opinia ta actuală. Nu există răspunsuri corecte sau greșite. Completează o singură variantă pentru fiecare întrebare, dacă nu este indicat altfel.",
        "",
    ]
    for question in questions:
        qid = str(question.get("question_id"))
        construct = str(question.get("construct") or "")
        prompt = str(question.get("prompt") or "").strip()
        response_type = str(question.get("response_type") or "")
        if not qid or not prompt or not response_type:
            raise ResearchPackError(f"question missing id/prompt/response type: {qid}")
        lines.extend([
            f"## {qid} — {construct}",
            "",
            prompt,
            "",
            _response_instruction(response_type),
            "",
            "Răspuns: ______________________________",
            "",
        ])
    return "\n".join(lines).rstrip() + "\n"


def questionnaire_docx(markdown_text: str, output_path: Path) -> None:
    document = Document()
    _set_document_defaults(document, title="Chestionar analiza de nevoi")
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            paragraph = document.add_paragraph(style="Title")
            paragraph.add_run(stripped[2:])
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=2)
        else:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(stripped)
            run.font.name = "Arial"
            run.font.size = Pt(10.5)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        temp = Path(handle.name)
    try:
        document.save(temp)
        _normalize_zip(temp, output_path)
    finally:
        temp.unlink(missing_ok=True)


def _write_csv(path: Path, headers, rows=()) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle, lineterminator="\n")
        writer.writerow(headers)
        for row in rows:
            writer.writerow(row)


def export_primary_research_pack(
    plan: Mapping[str, Any],
    output_dir: Path,
    *,
    project_id: str,
) -> Dict[str, Any]:
    questions = list(plan.get("questions") or [])
    if not questions:
        raise ResearchPackError("primary research plan has no questions")
    output_dir.mkdir(parents=True, exist_ok=True)

    questionnaire_md = output_dir / "CHESTIONAR_ELEVI.md"
    questionnaire_doc = output_dir / "CHESTIONAR_ELEVI.docx"
    population_csv = output_dir / "POPULATION_SNAPSHOT_TEMPLATE.csv"
    population_meta = output_dir / "POPULATION_SNAPSHOT_METADATA.json"
    raw_csv = output_dir / "RAW_RESPONSES_TEMPLATE.csv"
    instructions = output_dir / "README_RESEARCH.md"
    manifest_path = output_dir / "RESEARCH_PACK_MANIFEST.json"
    package_zip = output_dir / "PRIMARY_RESEARCH_PACK.zip"

    markdown = render_questionnaire_markdown(plan)
    _write_text(questionnaire_md, markdown)
    questionnaire_docx(markdown, questionnaire_doc)
    _write_csv(population_csv, ["grade", "qualification", "count"])
    _write_csv(raw_csv, ["respondent_id", "population_snapshot_id", "grade", "qualification", "question_id", "value"])

    metadata = {
        "schema_version": "nf.population_snapshot.template.v0.1",
        "project_id": project_id,
        "snapshot_id": "",
        "as_of_date": "",
        "school_year": "",
        "school_identity": "",
        "eligible_population_n": None,
        "grades_in_scope": [],
        "qualifications_in_scope": [],
        "source_document_id": "",
        "source_date": "",
        "source_hash_or_receipt": "",
        "instruction": "Completează metadata numai dintr-o evidență administrativă autorizată. Suma rândurilor din POPULATION_SNAPSHOT_TEMPLATE.csv trebuie să fie egală cu eligible_population_n.",
    }
    _write_text(population_meta, json.dumps(metadata, ensure_ascii=False, indent=2, sort_keys=True) + "\n")

    qids = [str(q["question_id"]) for q in questions]
    instruction_text = f"""# Instrucțiuni cercetare primară — Needs Factory

**Proiect:** {project_id}

## Ordinea obligatorie

1. Completează și validează mai întâi populația eligibilă pe clase și calificări.
2. Blochează `snapshot_id` înainte de aplicarea chestionarului.
3. Atribuie fiecărui elev un `respondent_id` pseudonimizat. Nu introduce nume, CNP, telefon sau e-mail în fișierul brut.
4. Aplică întrebările exact cu ID-urile: {', '.join(qids)}.
5. Introdu fiecare răspuns ca un rând separat în `RAW_RESPONSES_TEMPLATE.csv`.
6. `grade` și `qualification` trebuie să corespundă exact unei straturi din snapshotul de populație.
7. Nu calcula manual concluziile și nu modifica valorile pentru a susține activitățile proiectului. Needs Factory validează și agregă răspunsurile.

## Reguli de blocare

- fără snapshot autorizat: cercetarea rămâne blocată;
- numărul de respondenți nu poate depăși populația stratului;
- același respondent nu poate schimba clasa/calificarea între răspunsuri;
- întrebările sau scalele nu se reformulează după începerea colectării;
- datele sintetice sau exemplele nu pot fi promovate ca evidence de proiect.
"""
    _write_text(instructions, instruction_text)

    files = [questionnaire_md, questionnaire_doc, population_csv, population_meta, raw_csv, instructions]
    manifest = {
        "schema_version": "nf.primary_research_pack.v0.1",
        "project_id": project_id,
        "sampling_strategy": plan.get("sampling_strategy"),
        "population_snapshot_id": plan.get("population_snapshot_id"),
        "question_ids": qids,
        "question_count": len(qids),
        "files": {path.name: {"sha256": sha256_file(path)} for path in files},
        "policy": {
            "no_direct_identifiers": True,
            "population_snapshot_before_collection": True,
            "raw_data_required": True,
            "synthetic_data_forbidden_as_project_evidence": True,
        },
    }
    _write_text(manifest_path, json.dumps(manifest, ensure_ascii=False, indent=2, sort_keys=True) + "\n")
    files.append(manifest_path)
    _deterministic_zip([(path, path.name) for path in files], package_zip)
    manifest["package_zip"] = {"path": package_zip.name, "sha256": sha256_file(package_zip)}
    return manifest
